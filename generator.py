import json
import re
import docx
from google import genai
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ======================
#  CONFIGURAÇÃO GERAL
# ======================

API_KEY = "AIzaSyD8Pkkj62UhJgC8r8rzoE2NF3eV2CYyNZY"  # <<< COLOQUE SUA CHAVE AQUI
client = genai.Client(api_key=API_KEY)

REQUISITOS_DOCX = "requisitos.docx"
SAIDA_JSON = "saida.json"
TESTES_JSON = "testes.json"
WORD_OUTPUT = "cenarios_de_testes.docx"

print("🚀 Gerador de cenários de testes iniciado.")

# =============================================
#  1. EXTRAÇÃO DO .DOCX → JSON DE REQUISITOS
# =============================================

def extrair_requisitos_docx(caminho=REQUISITOS_DOCX):
    print("📄 Extraindo requisitos do DOCX")

    doc = docx.Document(caminho)

    sections = []
    sec_atual = {"title": "", "requirements": []}

    for para in doc.paragraphs:
        texto = para.text.strip()
        if not texto:
            continue

        # Títulos são negrito
        if para.runs and any(r.bold for r in para.runs):
            if sec_atual["title"]:
                sections.append(sec_atual)
            sec_atual = {"title": texto, "requirements": []}
        else:
            sec_atual["requirements"].append({"text": texto})

    if sec_atual["title"]:
        sections.append(sec_atual)

    # salvar JSON
    with open(SAIDA_JSON, "w", encoding="utf-8") as f:
        json.dump({"sections": sections}, f, ensure_ascii=False, indent=2)

    print("✅ Requisitos extraídos e salvos em no arquivo 'saida.json'")

    return sections

# =================================================================================================
#  2. PROMP PARA GERAÇÃO DOS CENÁRIOS (PODE SER AJUSTADO E MELHORADO CONFORME A SUA NECESSIDADE)
# =================================================================================================

QA_PROMPT = """
Você é um Engenheiro de QA Sênior com 15+ anos de experiência em testes manuais, automatizados, análise de requisitos, modelagem de cenários e testes baseados em risco.  
Sua missão é gerar uma suíte de testes **completa, técnica, rastreável e pronta para execução**, baseada exclusivamente nos requisitos fornecidos.

# 🎯 OBJETIVO
Produzir uma suíte de testes completa, robusta e tecnicamente abrangente, contemplando cenários funcionais positivos, negativos, casos de exceção, limites mínimos e máximos de entrada (tamanho, tipo, caracteres especiais, números), além de validações de segurança e performance. O objetivo final é garantir cobertura total dos requisitos, detecção antecipada de falhas ocultas e zero ambiguidade em cada cenário descrito.

# 🔎 METODOLOGIA

## 1. ANÁLISE ESTRUTURAL
Extraia de forma explícita:
- Entidades principais, atributos e relacionamentos
- Regras de negócio essenciais e condicionais
- Fluxos primários, alternativos e exceções
- Dependências externas e integrações
- Riscos técnicos, funcionais e de usabilidade

## 2. TÉCNICAS DE TESTE OBRIGATÓRIAS
Utilize e informe quais técnicas sustentam cada cenário:
- Particionamento de Equivalência
- Análise de Valor Limite
- Tabela de Decisão
- Testes Baseados em Estado
- Testes Exploratórios e Heurísticas (SFDPOT, HICCUPPS)
- Análise de Risco

## 3. COBERTURA MÍNIMA NECESSÁRIA
Cada suíte deve contemplar:
- Happy path completo
- Validações de dados (tipo, formato, tamanho, regex, range)
- Permissões, níveis de acesso e autenticação
- Comportamentos inesperados, erros e exceções
- Performance (SLAs definidos ou padrão: < 2s para 95% das requisições)
- Segurança (OWASP Top 10 + autenticação/autorizações incorretas)
- Compatibilidade cross-browser e cross-device
- Persistência e integridade de dados
- Cenários assíncronos e concorrência (quando aplicável)

# 📦 FORMATO DE ENTREGA
Retorne **somente JSON válido**, sem markdown, sem textos extras.

Estrutura padrão obrigatória:

{
  "analise_requisitos": {
    "entidades": [],
    "atributos_criticos": [],
    "regras_negocio": [],
    "fluxos": {
      "principal": [],
      "alternativos": [],
      "excecoes": []
    },
    "integracoes": [],
    "riscos": []
  },

  "cenarios_funcionais": [
    {
      "id": "TC-FUNC-001",
      "titulo": "",
      "categoria": "CRUD|Fluxo|RegraNegocio|Integracao",
      "prioridade": "Crítica|Alta|Média|Baixa",
      "tecnica_teste": "",
      "descricao": "",
      "pre_condicoes": [],
      "dados_teste": {},
      "passos": [],
      "resultado_esperado": "",
      "criterios_aceitacao": [],
      "pos_condicoes": ""
    }
  ],

  "cenarios_negativos": [],
  "cenarios_borda": [],
  "cenarios_seguranca": [],
  "cenario_performance": [],
  "bugs_provaveis": [],
  "matriz_rastreabilidade": [],
  "metricas_qualidade": {
    "cobertura_requisitos": "",
    "total_casos_teste": 0,
    "distribuicao_por_categoria": {}
  }
}

# ⚠️ REGRAS CRÍTICAS E INEGOCIÁVEIS

1. IDs devem seguir: TC-{CATEGORIA}-{NNN}
2. Nenhum passo pode ser vago — todos devem ser acionáveis
3. Resultados devem ser 100% mensuráveis e verificáveis
4. Testes devem considerar condições de concorrência sempre que possível
5. Nunca incluir textos fora do JSON, nem comentários
6. Nada de vírgulas sobrando (JSON deve ser validado mentalmente por um ninja)
7. Sempre mapear pelo menos 1 bug provável por regra de negócio

# 🏆 EXEMPLO DO QUE ESPERO
❌ Vago: "Testar login"
✅ Robusto: "Login com credenciais válidas deve retornar token JWT, registrar timestamp do login e responder em < 2s"

Retorne APENAS o JSON, sem texto adicional."""

def build_prompt():
    print("🧩 Pensando para criar os melhores cenários")

    with open(SAIDA_JSON, "r", encoding="utf-8") as f:
        requisitos = json.load(f)

    combined = ""
    for section in requisitos["sections"]:
        combined += f"\nSEÇÃO: {section['title']}\n"
        for req in section["requirements"]:
            combined += f"- {req['text']}\n"

    return QA_PROMPT + "\n\nREQUISITOS ANALISADOS:\n" + combined

# ==================================
#  3. LIMPAR JSON VINDO DO GEMINI
# ==================================

def limpar_json_bruto(texto):
    try:
        match = re.search(r'\{.*\}', texto, re.DOTALL)
        return match.group(0) if match else texto
    except:
        return texto

# ====================
#  4. CHAMAR GEMINI
# ====================

def gerar_cenarios(prompt):
    print("🤖 Gemini está Processando as informações.")
    resp = client.models.generate_content(
        model="models/gemini-2.5-flash", #caso queira utilizar outro modelo do gemini basta trocar por outro. Ex: "gemini-2.0-flash-lite"
        contents=prompt
    )
    return resp.text

# ===========================================================================================
#  5. FORMATAR CÉLULAS DO WORD (VERSÃO INICIAL DO TEMPLATE DO WORD AINDA PODE SER MELHORADO)
# ===========================================================================================

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)

def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")

    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:color"), "808080")
        borders.append(el)

    tcPr.append(borders)

def style_header(cell):
    set_cell_bg(cell, "D9D9D9")
    set_cell_borders(cell)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(10)

# =====================================================================
#  6. GERAR TEMPLATE DO WORD PREENCHIDO COM DADOS ANALISADO PELO GEMINI
# =====================================================================

def salvar_word(json_data):
    print("📝 Gerando seu plano de testes em Word")

    doc = docx.Document()
    doc.add_heading("Cenários de Teste - IA Generator", level=1)

    def add_table(title, itens):
        doc.add_heading(title, level=2)
        table = doc.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells

        headers = ["ID", "Título", "Descrição", "Pré-condições", "Passos", "Resultado Esperado"]

        for i, h in enumerate(headers):
            hdr[i].text = h
            style_header(hdr[i])

        for item in itens:
            row = table.add_row().cells
            row[0].text = item.get("id", "")
            row[1].text = item.get("titulo", "")
            row[2].text = item.get("descricao", "")
            row[3].text = "\n".join(item.get("pre_condicoes", []))
            row[4].text = "\n".join(item.get("passos", []))
            row[5].text = item.get("resultado_esperado", "")

            for c in row:
                set_cell_borders(c)

        doc.add_paragraph("")

    if "cenarios_funcionais" in json_data:
        add_table("Cenários Funcionais", json_data["cenarios_funcionais"])

    if "cenarios_negativos" in json_data:
        add_table("Cenários Negativos", json_data["cenarios_negativos"])

    if "cenarios_borda" in json_data:
        add_table("Cenários de Borda", json_data["cenarios_borda"])

    doc.save(WORD_OUTPUT)
    print(f"Seu documento foi gerado ✅. Confira o arquivo: '{WORD_OUTPUT}'")

# ============================================================
#  7. EXECUÇÃO PRINCIPAL
# ============================================================

if __name__ == "__main__":
    extrair_requisitos_docx()

    prompt = build_prompt()
    resposta = gerar_cenarios(prompt)

    resposta_limpa = limpar_json_bruto(resposta)

    with open(TESTES_JSON, "w", encoding="utf-8") as f:
        f.write(resposta_limpa)

    try:
        json_data = json.loads(resposta_limpa)
    except:
        print("Gemini retornou um JSON inválido. Isso pode ser um erro :( ❌")
        print("JSON bruto salvo em 'testes.json'")
        exit()

    salvar_word(json_data)

    print("🎉 Plano de testes criado com sucesso!")